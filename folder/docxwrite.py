# -*- coding:utf-8 -*-

import os

from docx import Document
from docx.shared import Inches, Pt

from datetime import datetime, timedelta

def docxwrite(dic):
	

	data = dic

	document = Document()
	
	document.add_picture ('{}/XDF.jpg'.format(os.path.dirname(__file__)), width = Inches(1.25))

	document.add_heading(u'                    优惠活动申请单', 0)

	document.add_paragraph('\n')

	paragraph = document.add_paragraph()
	run = paragraph.add_run(u'申请人:{}     所在部门:{}     申请日期:{}'.format(data["recordername"], data["departmentname"], data["recordtime"]))
	run.font.size = Pt(12)
	run.font.bold = False
	run.font.name = u'黑体'

	table0 = document.add_table(rows = 1, cols = 4)

	fir_cells = table0.rows[0].cells

	fir_cells[0].text, fir_cells[1].text, fir_cells[2].text, fir_cells[3].text = u'优惠编码', data["coreid"], u'优惠名称', data["couponname"]


	table = document.add_table(rows = 4, cols = 2)



	cells0, cells1, cells2, cells3 = table.rows[0].cells, table.rows[1].cells, table.rows[2].cells, table.rows[3].cells

	cells0[0].width = 1

	cells0[1].width = 1

	cells0[0].text, cells0[1].text = u'优惠活动具体内容', data["content"]



	cells1[0].text, cells1[1].text = u'优惠时间', '{}--{}'.format(data["starttime"], data["endtime"])
	

	cells2[0].text, cells2[1].text = u'优惠原因', data["advice"]


	cells3[0].text, cells3[1].text = u'优惠金额(￥)/折扣', data["account"]



	table2 = document.add_table(rows = 5, cols = 2)

	checktimelist = data["checktimelist"].replace('[', '').replace(']', '').replace('\'','').split(',')

	ano_cell0, ano_cell1, ano_cell2, ano_cell3, ano_cell4 = table2.rows[0].cells, table2.rows[1].cells, table2.rows[2].cells, table2.rows[3].cells, table2.rows[4].cells

	ano_cell0[0].text, ano_cell0[1].text = u'申请人签字:               {}'.format(data["recordername"]), u'日期:{}'.format(datetime.strptime(checktimelist[0], '%Y-%m-%d %H:%M:%S'))

	ano_cell1[0].text, ano_cell1[1].text = u'部门主管签字:           {}'.format(data["leadername"]), u'日期:{}'.format(datetime.strptime(checktimelist[1][:-8], '%Y-%m-%d %H:%M:%S'))

	ano_cell2[0].text, ano_cell2[1].text = u'市场部主管签字:       龙泉',  u'日期:{}'.format(datetime.strptime(checktimelist[2][:-8], '%Y-%m-%d %H:%M:%S'))

	ano_cell3[0].text, ano_cell3[1].text = u'教务部门主管签字:  吴华强', u'日期:{}'.format(datetime.strptime(checktimelist[3][:-8], '%Y-%m-%d %H:%M:%S'))
 
	ano_cell4[0].text, ano_cell4[1].text = u'校长签字:                     刘清文', u'日期:{}'.format(datetime.strptime(checktimelist[4][:-8], '%Y-%m-%d %H:%M:%S'))

	document.add_paragraph(
		u'\n\n备注: 优惠活动必须填写上面的申请单, 如无校长签名, 须有校长短信通知教务部, 任何部门如不按以上要求操作, 教务部有权不提供公章和基于优惠报名, 特此通知。\n\n'
	)	


	document.add_paragraph(
		u'                                                                                                                                                成都新东方学校'
	)
	document.add_paragraph(
		u'                                                                                                                                                     教务部'
	)		



        filename = 'coupon_docx_{0}.docx'.format(datetime.now().strftime("_%b_%d_%Y_%H_%M_%S_"))
        filefullpath = '{}/{}'.format(os.path.dirname(__file__), filename)
        document.save(filefullpath)
        return '/download/docx/{}'.format(filename)	

